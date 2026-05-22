import io
import logging
from pathlib import Path

import fitz
from PIL import Image

from app.core.config import settings
from app.services.preprocessing import normalize_sanskrit_text, remove_repeated_headers_footers

logger = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def load_pdf(file_path: str | Path) -> list[dict]:
    """Extract text page-by-page from a PDF, with optional OCR fallback."""
    path = Path(file_path)
    pages = []

    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc):
            text = page.get_text("text")
            extraction_mode = "native"

            if settings.enable_ocr and len(text.strip()) < settings.ocr_min_chars:
                try:
                    import pytesseract

                    pix = page.get_pixmap(dpi=300)
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    text = pytesseract.image_to_string(image, lang="san+eng")
                    extraction_mode = "ocr"
                except Exception as exc:
                    logger.warning("OCR failed for %s page %s: %s", path.name, page_index + 1, exc)

            text = normalize_sanskrit_text(text)
            if text:
                pages.append(
                    {
                        "filename": path.name,
                        "page_number": page_index + 1,
                        "text": text,
                        "source_type": "pdf",
                        "extraction_mode": extraction_mode,
                    }
                )

    return remove_repeated_headers_footers(pages)


def load_txt(file_path: str | Path) -> list[dict]:
    path = Path(file_path)
    raw = path.read_text(encoding="utf-8", errors="replace")
    text = normalize_sanskrit_text(raw)
    if not text:
        return []
    return [
        {
            "filename": path.name,
            "page_number": 1,
            "text": text,
            "source_type": "txt",
            "extraction_mode": "txt",
        }
    ]


def load_docx(file_path: str | Path) -> list[dict]:
    path = Path(file_path)
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to ingest .docx files. "
            "Install backend dependencies with: pip install -r backend/requirements.txt"
        ) from exc

    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    table_cells = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_cells.append(" | ".join(cells))

    text = normalize_sanskrit_text("\n".join(paragraphs + table_cells))
    if not text:
        return []
    return [
        {
            "filename": path.name,
            "page_number": 1,
            "text": text,
            "source_type": "docx",
            "extraction_mode": "docx",
        }
    ]


def load_document(file_path: str | Path) -> list[dict]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".txt":
        return load_txt(path)
    if suffix == ".docx":
        return load_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def iter_supported_documents(folder: str | Path) -> list[Path]:
    path = Path(folder)
    return sorted(p for p in path.iterdir() if p.suffix.lower() in SUPPORTED_EXTENSIONS and p.is_file())
